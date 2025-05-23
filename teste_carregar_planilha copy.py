from tkinter import *
from tkinter import ttk
from tkinter import filedialog, messagebox
import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from docx.shared import Inches
from docxtpl import DocxTemplate, InlineImage
import requests
from io import BytesIO
from docxtpl import InlineImage

global context
context = {}
#doc = Document()

caminho_template = "C:/Users/walyson.ferreira/Desktop/openpy/template/TEMPLATE DEFINITIVO.docx"

arquivo = load_workbook('OpenPy.xlsx')
planilha = arquivo.active

def adicionar_campo(cn, v):
    
    context
    valor = v
    coluna = cn

    nome_documento = entrada_nome_arquivo.get().strip()
    if not nome_documento:
        nome_documento = 'teste_funcao.docx'
    else:
        nome_documento += '.docx'
        
    caminho_template = "C:/Users/walyson.ferreira/Desktop/openpy/template/TEMPLATE DEFINITIVO.docx"
    
    
    # nome_documento
    doc = Document('AAAAAAAA.docx')
    
    if coluna != "":
        coluna = str(coluna)
        coluna = coluna.lower()
        coluna = coluna.replace(" ", "_")
  
    # Adicionando valores placeholders
    doc.add_heading(cn + ': {{' + coluna + '}}', level=7)
    
    doc.save('AAAAAAAA.docx')
    
    print(f"Nome da Coluna: {coluna}", f"Valor de CN: {cn}",f"\nNome do Valor: {valor}", f"\nAdicionado ao arquivo!{doc}")
    
    doc_template = DocxTemplate('AAAAAAAA.docx')
    
    context[str(coluna)] = valor
    #context["nome_projeto"] = entrada_nome_arquivo.get().strip()
    # VALIDAÇÃO
    print(f"Nome do arquivo: {nome_documento} setado com sucesso!")
    # Condição para nome de coluna??
    doc_template.render(context)
    doc_template.save('AAAAAAAA.docx')
    
    
    
    
    
    # VALIDAÇÃO
    print(f"Nome do arquivo: {nome_documento} setado com sucesso!")
    # Condição para nome de coluna??
    # doc.render(context)
    
    

    

def selecionar_template():
    
    global caminho_template
    
    arquivo = filedialog.askopenfilename()
    #arquivo = "C:/Users/walyson.ferreira/Desktop/openpy/template/TEMPLATE DEFINITIVO.docx"
    
    
    
    print(f"CAMINHO DO ARQUIVO SELECIONADO: {arquivo}")
    
    caminho_template = arquivo
    
    
# def substituir_placeholder_por_imagem():
#     try:
#         docx_path = entrada_nome_arquivo.get().strip() + ".docx"
#         doc = Document(docx_path)
#         imagem_path = filedialog.askopenfilename()

        

#         for paragraph in doc.paragraphs:
#             if "{{ imagem }}" in paragraph.text:
#                 # Remove apenas o texto do placeholder
#                 paragraph.clear()  # Limpa o conteúdo do parágrafo, mas mantém a posição

#                 # Adiciona a imagem dentro do mesmo parágrafo
#                 run = paragraph.add_run()
#                 run.add_picture(imagem_path, width=Inches(1))
#                 break


#         output_path = os.path.join(os.path.dirname(entrada_nome_arquivo.get().strip()), docx_path)
#         doc.save(output_path)
#         messagebox.showinfo("Sucesso", f"Arquivo salvo como:\n{output_path}")

#     except Exception as e:
#         messagebox.showerror("Erro", f"Ocorreu um erro:\n{str(e)}")



# def substituir_placeholder_por_imagem():
#     try:
#         caminho_arquivo = entrada_nome_arquivo.get().strip() + ".docx"
#         imagem_path = filedialog.askopenfilename(filetypes=[("Imagens", "*.png *.jpg *.jpeg *.bmp")])
#         if not imagem_path:
#             return

#         # Usando DocxTemplate e InlineImage para substituir o placeholder
#         doc = DocxTemplate(caminho_arquivo)

#         # Adiciona a imagem ao contexto para o placeholder {{ imagem }}
#         context["imagem"] = InlineImage(doc, imagem_path, width=Inches(2))

#         doc.render(context)
#         doc.save(caminho_arquivo)
#         messagebox.showinfo("Sucesso", f"Imagem inserida e arquivo salvo como:\n{caminho_arquivo}")
#     except Exception as e:
#         messagebox.showerror("Erro", f"Ocorreu um erro:\n{str(e)}")
        
        
def selecionar_imagem(self):
    self.imagem_path = filedialog.askopenfilename(filetypes=[("Imagens", "*.png *.jpg *.jpeg")])
    if self.imagem_path:
        messagebox.showinfo("Imagem selecionada", f"Imagem: {self.imagem_path}")    

def adicionar_imagem():

    nome_arquivo = entrada_nome_arquivo.get().strip()
    image_path = filedialog.askopenfilename()

    if not nome_arquivo:
        nome_arquivo =  "default.docx"
    else:
        nome_arquivo += ".docx"
    
    try:
        doc = Document(nome_arquivo)
        print(f"Abrindo o arquivo existente. {nome_arquivo}")
    except:
        print(f"Arquivo '{nome_arquivo}' não encontrado. Criando um novo arquivo")
        doc = Document()

    largura_em_polegadas = 5.5
    # Adicionando imagem
    #doc.add_picture(image_path, width=Inches(2), height=Inches(2))
    
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{ imagem }}" in paragraph.text:
            # Remove o parágrafo com o placeholder
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

            # Inserir a imagem no mesmo lugar
            novo_paragrafo = doc.paragraphs.insert(i, doc.add_paragraph())
            run = novo_paragrafo.add_run()
            run.add_picture(image_path, width=Inches(largura_em_polegadas))
            print("Imagem adicionada com sucesso!")
            break
        
        
    doc.save(nome_arquivo)
    print("arquivo salvo com sucesso!")

def adicionar_cabecalho(variavel):
    # Get the filename from the Entry widget
    nome_arquivo = entrada_nome_arquivo.get().strip()
    

    if not nome_arquivo:  # Handle empty input
        print("O campo de nome do arquivo está vazio. Usando 'default.docx' como padrão.")
        nome_arquivo = "default.docx"
    else:
        nome_arquivo += ".docx"  # Add the .docx extension

    try:
        # Try to open the specified Word file
        doc = Document(nome_arquivo)
        print(f"Abrindo o arquivo existente: {nome_arquivo}")
    except Exception as e:
        # If the file doesn't exist, create a new document
        print(f"Arquivo '{nome_arquivo}' não encontrado. Criando um novo arquivo.")
        doc = Document()

    # Add the heading to the document
    doc.add_heading(variavel, level=1)

    doc.add_section

    doc.save(nome_arquivo)
    print(f"Cabeçalho adicionado ao arquivo: {nome_arquivo}")


def adicionar_paragrafo(variavel):
    # Get the filename from the Entry widget
    nome_arquivo = entrada_nome_arquivo.get().strip()
    
    if not nome_arquivo:  # Handle empty input
        print("O campo de nome do arquivo está vazio. Usando 'default.docx' como padrão.")
        nome_arquivo = "default.docx"
    else:
        nome_arquivo += ".docx"  # Add the .docx extension

    try:
        # Try to open the specified Word file
        doc = Document(nome_arquivo)
        print(f"Abrindo o arquivo existente: {nome_arquivo}")
    except Exception as e:
        # If the file doesn't exist, create a new document
        print(f"Arquivo '{nome_arquivo}' não encontrado. Criando um novo arquivo.")
        doc = Document()

    # Add the paragraph to the document
    doc.add_paragraph(variavel)
    
    if not nome_arquivo:
        context["nome_projeto"] = nome_arquivo
  
    doc.save(nome_arquivo)
    print(f"Parágrafo adicionado ao arquivo: {nome_arquivo}")



def pesquisa():
    global df, tamanho_lista  # Ensure df and tamanho_lista are accessible
    search_text = entrada_texto.get().strip()

    if not search_text:
        print("Campo de pesquisa vazio.")
        return

    if 'df' in globals():
        # Filter rows where any cell contains the search text (case-insensitive)
        filtered_df = df[df.apply(lambda row: row.astype(str).str.contains(search_text, case=False).any(), axis=1)]

        # Update tamanho_lista with the number of filtered rows
        tamanho_lista = len(filtered_df)
        resultado_texto.set(f"Resultados encontrados: {tamanho_lista}") # Update the label text

        # Clear the Treeview
        for col in tree.get_children():
            tree.delete(col)

        # Insert filtered data into the Treeview
        for index, row in filtered_df.iterrows():
            tree.insert("", "end", values=list(row))

        print(f"Exibindo resultados para: {search_text}")
    else:
        print("Nenhum arquivo carregado para pesquisa.")


def open_new_window():
    
    # if caminho_template:
    #     print(f"caminho de template selecionado! valor: {caminho_template}")
    # else:
    #     print("Nenhum template selecionado!")
    
    selected_item = tree.focus()  # Get selected item ID
    if selected_item:
        values = tree.item(selected_item, 'values')  # Get row data

        # Create new window
        new_win = Toplevel(root)
        new_win.title("Detalhes")

        # Get screen dimensions
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()

        # Calculate 1/4 of the screen size
        window_width = screen_width // 2
        window_height = screen_height // 2

        # Position the window in the right-top corner
        x = screen_width - window_width
        y = 0

        # Set the geometry of the new window
        new_win.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Create a frame for the canvas and scrollbar
        frame = Frame(new_win)
        frame.pack(fill=BOTH, expand=True)

        # Create a canvas inside the frame
        canvas = Canvas(frame)
        canvas.pack(side=LEFT, fill=BOTH, expand=True)

        # Add a scrollbar to the canvas
        scrollbar = ttk.Scrollbar(frame, orient=VERTICAL, command=canvas.yview)
        scrollbar.pack(side=RIGHT, fill=Y)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        # Create another frame inside the canvas for the content
        content_frame = Frame(canvas)
        canvas.create_window((0, 0), window=content_frame, anchor='nw')

        column_names = tree["columns"]

        interruptor = False
        
            
        def salvar_todos():
            nonlocal interruptor
            interruptor != interruptor
            print("Botão clicado!")

        # Button(content_frame, text="Salvar TODOS", command=salvar_todos).grid(row=1, column=3, padx=10, pady=5)
        # Add the labels to the content frame
        for i, val in enumerate(values):
                
                column_name = column_names[i]

                # Column name label
                Label(content_frame, text=column_name, font=('Arial', 10, 'bold')).grid(row=i, column=0, sticky='w', padx=10, pady=5)

                # Selectable value using readonly Entry
                entry = ttk.Entry(content_frame, font=('Arial', 10), width=30)
                entry.insert(0, val)
                entry.config(state='readonly')
                entry.grid(row=i, column=1, sticky='w', padx=10, pady=5)
                
                Button(content_frame, text="Salvar TODOS", command=salvar_todas_colunas).grid(row=1, column=4, padx=10, pady=5)
                # Always run this block when the button is clicked
                # if interruptor:
                #     column_name = column_names[i]
                
                #     nome_documento = 'teste_funcao.docx'
                #     caminho_template = 'C:\\Users\\walyson.ferreira\\Desktop\\openpy\\template\\template.docx'
                    
                #     doc = DocxTemplate(caminho_template)
                    
                #     if 'context' not in globals():
                #         global context
                #         context = {}
                    
                #     # CONVERTE NOME DE COLUNAS PARA MINÚSCULO E TROCAR ESPAÇOS " " POR "_" 
                #     print(f"\033[1mVALOR ANTES DA CONVERSÃO DE CN: {column_name}\033[0m")
                #     if column_name != "":
                #         column_name = str(column_name)
                #         column_name = column_name.lower()
                #         column_name = column_name.replace(" ", "_")
                #     print(f"\033[1mVALOR ATUAL DE CN: {column_name}\033[0m")
                    
                #     context[str(column_name)] = val
                #     doc.render(context)
                #     doc.save(nome_documento)
                #     print(f"Nome da Coluna: {column_name}", f"\nNome do Valor: {val}", "\nAdicionado ao arquivo!")
                    
                #     interruptor = False
                
                
                
                # *FUNÇÃO CORINGA*
                if column_names[i] == column_names[i]: 
                    Button(content_frame, text="Adicionar dado ao relatório", command=lambda v=val, cn=column_name: adicionar_coluna(cn, v)).grid(row=i, column=2, padx=10, pady=5)
                    Button(content_frame, text="Adicionar campo ao relatório", command=lambda valor=val, coluna=column_name: adicionar_campo(coluna, valor)).grid(row=i, column=3, padx=10, pady=5)
                    
              
                
        # Update scrollregion when contents are added
        content_frame.update_idletasks()
        canvas.config(scrollregion=canvas.bbox("all"))

        Button(content_frame,
                    text="Salvar Relatório",
                    command=salvar_todos
                    ).grid(row=i, column=2, padx=10, pady=5)
    else:
        print("Nenhuma linha selecionada.")


# FUNÇÃO SALVAR TODOS
def salvar_todas_colunas():
    # declaração de contexto
    global context
    # CAMINHO DA IMAGEM
    #caminho_imagem = filedialog.askopenfilename()
    
    
    selected_item = tree.focus()
    
    if not selected_item:
        print("Nenhuma linha selecionada.")
        return

    values = tree.item(selected_item, 'values')
    column_names = tree["columns"]


    
    
    nome_documento = entrada_nome_arquivo.get().strip()
    titulo_documento = entrada_nome_arquivo.get().strip()
    
    if nome_documento == "":
        nome_documento = 'teste_funcao.docx'
    else:
        nome_documento += ".docx"

    # SELECIONAR ARQUIVO PARA PEGAR CAMINHO DO TEMPLATE
    #caminho_template = 'C:\\Users\\walyson.ferreira\\Desktop\\openpy\\template\\template.docx'
    
    from docxtpl import DocxTemplate  # Ensure import

    #caminho_template = "C:/Users/walyson.ferreira/Desktop/openpy/template/TEMPLATE DEFINITIVO.docx"
    
    
    print(f"Caminho do template: {caminho_template}")
    doc = DocxTemplate(caminho_template)
   
    if 'context' not in globals():
        context = {}

    for i, val in enumerate(values):
        cn = column_names[i]
        print(f"\033[1mVALOR ANTES DA CONVERSÃO DE CN: {cn}\033[0m")
        if cn != "":
            cn = str(cn).lower().replace(" ", "_")
            cn = str(cn).replace("ç", "c")
            cn = str(cn).replace("ã", "a")
        print(f"\033[1mVALOR ATUAL DE CN: {cn}\033[0m")
        
        # necessário para setar os valores das colunas
        teste = cn
        teste = cn.strip()
        
        if teste == "obra_parada/em_andamento":
            teste = "obra_parada"
            cn = teste
        elif teste == "%_execução":
            teste = "execucao"
            cn = teste
        
            
        context[cn] = val
        print(f"Valor atual da coluna: {cn}", f"\nValor atual do valor: {val}")
        # setando valor do título do documento
        context["nome_projeto"] = titulo_documento
        
    # Adicionando ao contexto o valor da variável IMAGEM
    #context["imagem"] = InlineImage(doc, caminho_imagem, width=Inches(2))
    # Renderizar contexto
    doc.render(context)
    doc.save(nome_documento)
    print("Todos os valores da linha foram adicionados ao arquivo!")

# Add this button in your open_new_window function, after the content_frame is created:
#Button(content_frame, text="Salvar TODOS", command=salvar_todas_colunas).grid(row=0, column=4, padx=10, pady=5)

# *FUNÇÃO CORINGA*
def adicionar_coluna(cn, v):
    
    # if 'context' not in globals():
    context
    
    valor = v
    coluna = cn

    nome_documento = entrada_nome_arquivo.get().strip()
    if not nome_documento:
        nome_documento = 'teste_funcao.docx'
    else:
        nome_documento += '.docx'
        
    #caminho_template = 'C:\\Users\\walyson.ferreira\\Desktop\\openpy\\template\\template.docx'
    
    
    doc = DocxTemplate(caminho_template)
    
    # CONVERTE NOME DE COLUNAS PARA MINÚSCULO E TROCAR ESPAÇOS " " POR "_" 
    #print(f"\033[1mVALOR ANTES DA CONVERSÃO DE CN: {coluna}\033[0m")
    
    
    
    if coluna != "":
        coluna = str(coluna)
        coluna = coluna.lower()
        coluna = coluna.replace(" ", "_")
    # Print with ANSI escape code for bold in supported terminals
    #print(f"\033[1mVALOR ATUAL DE CN: {coluna}\033[0m")


    context[str(coluna)] = valor
    context["nome_projeto"] = entrada_nome_arquivo.get().strip()
    
    # VALIDAÇÃO
    print(f"Nome do arquivo: {nome_documento} setado com sucesso!")
    # Condição para nome de coluna??
    
    
    doc.render(context)
    doc.save(nome_documento)
    print(f"Nome da Coluna: {coluna}", f"\nNome do Valor: {valor}", "\nAdicionado ao arquivo!")
    


def limpar_contexto():
    global context
    print(f"Valor do contexto ANTES de limpar: {context}")
    context.clear()
    print(f"Valor do contexto DEPOIS de limpar: {context}")
    #print("Contexto limpo com sucesso!")
    
    
def load_file():
    # print(f"template carregado: {caminho_template}")
    global df     
            
    caminho_arquivo = filedialog.askopenfilename()
    
    #caminho_arquivo = 'https://integracao-my.sharepoint.com/:x:/r/personal/walyson_ferreira_integracao_gov_br/_layouts/15/Doc.aspx?sourcedoc=%7B7B2FC40F-4C95-4949-BA1D-7DC02D8FF33F%7D&file=Pasta%201.xlsx'
    
    
    print(f"Valor do arquivo atual: {caminho_arquivo}")
    # Caso o arquivo seja carregado, verificar se é um arquivo de planilha
    if caminho_arquivo:
        nome_arquivo, extensao_arquivo = os.path.splitext(caminho_arquivo)

        if extensao_arquivo == '.csv':
            try:
                df = pd.read_csv(caminho_arquivo, encoding='utf-8', sep=';')
                print("Arquivo de planilha carregado")

            except UnicodeDecodeError:
                df = pd.read_csv(caminho_arquivo, encoding='ISO-8859-1', sep=';')
                print("Arquivo de planilha carregado")

        elif extensao_arquivo in ['.xlsx', '.xls']:
            df = pd.read_excel(caminho_arquivo)
            print("Arquivo de planilha carregado")
            
        else:   
            
            
            
            
            
            
            df = pd.read_excel(caminho_arquivo, engine="openpyxl")
            print("tipo de arquivo não permitido")
            return

        # limpar dados do treeview
        for col in tree.get_children():
            tree.delete(col)
        
        # definir colunas do treeview
        
        tree["columns"] = list(df.columns)
        tree["show"] = "headings" # oculta a primeira coluna vazia

        # configura os títulos das colunas
        for coluna in df.columns:
            tree.heading (coluna, text=coluna)
            tree.column(coluna, anchor='center', minwidth=100)
        
        # insere os dados no Treeview
        for index, row in df.iterrows():
            tree.insert("", "end", values=list(row))


# centralizar window
def center_window(width=300, height=200):
    # Get screen width and height
    screen_width = root.winfo_screenwidth() * 2
    screen_height = root.winfo_screenheight()

    # Set the window width to a fixed size (e.g., 1/4 of the screen width)
    window_width = screen_width // 4
    window_height = screen_height  # Full height of the screen

    # Position the window on the left side of the screen
    x = 0  # Left edge
    y = 0  # Top edge

    # Set the geometry of the root window
    root.geometry(f"{window_width}x{window_height}+{x}+{y}")


root = Tk()
root.title("Planilha")
center_window(500, 400)
root.resizable(True, True)

root.grid_rowconfigure(2, weight=1)
root.grid_columnconfigure(0, weight=1)
# Create a StringVar to hold the label text
resultado_texto = StringVar()
resultado_texto.set(f"Resultados encontrados: {0}")
# Initialize with 0

# Use grid for all root widgets
entrada_texto = ttk.Entry(root, width=20) 
entrada_texto.grid(row=0, column=0, padx=5, pady=10, sticky="w")

pesquisa_btn = ttk.Button(root, text="Pesquisar", command=pesquisa)
pesquisa_btn.grid(row=0, column=1, padx=5, pady=10, sticky="w")

detalhes_btn = ttk.Button(root, text="Detalhes", command=open_new_window)
detalhes_btn.grid(row=0, column=2, padx=5, pady=10, sticky="w")

carregar_btn = ttk.Button(root, text="Carregar Arquivo", command=load_file)
carregar_btn.grid(row=0, column=3, padx=5, pady=10, sticky="w")

# imagem_btn = ttk.Button(root, text="Carregar Imagem", command=substituir_placeholder_por_imagem)
# imagem_btn.grid(row=0, column=4, padx=5, pady=10, sticky="w")

# BOTÃO LIMPAR CONTEXTO
limpar_btn = ttk.Button(root, text="Limpar Contexto", command=limpar_contexto)
limpar_btn.grid(row=0, column=7, padx=5, pady=5)

# Botão adicionar TEMPLATE

add_template_btn = ttk.Button(root, text="Selecionar Template", command=selecionar_template)
add_template_btn.grid(row=0, column=8, padx=5, pady=5)


entrada_nome_arquivo = ttk.Entry(root, width=15)
entrada_nome_arquivo.grid(row=0, column=5, padx=5, pady=10, sticky="w")

# "Resultados encontrados" label in row=1
filtro_numero = ttk.Label(root, textvariable=resultado_texto)
filtro_numero.grid(row=1, column=0, columnspan=6, padx=5, pady=0 ,sticky="w")

# Use pack/grid for frame as before
frame = ttk.Frame(root)
frame.grid(row=2, column=0, columnspan=9, sticky="nsew", padx=10, pady=10)

# Scrollbars
scroll_y = ttk.Scrollbar(frame, orient="vertical")
scroll_x = ttk.Scrollbar(frame, orient="horizontal")

tree = ttk.Treeview(frame, yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

# Use grid instead of place
frame.columnconfigure(0, weight=1)
frame.rowconfigure(0, weight=1)

tree.grid(row=0, column=0, sticky="nsew")
scroll_y.grid(row=0, column=1, sticky="ns")
scroll_x.grid(row=1, column=0, sticky="ew")

scroll_y.config(command=tree.yview)
scroll_x.config(command=tree.xview)

root.mainloop()