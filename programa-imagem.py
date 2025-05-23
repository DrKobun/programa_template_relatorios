import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
import os

# Função principal de substituição do placeholder
def substituir_placeholder_por_imagem(docx_path, imagem_path):
    try:
        
        doc = Document(docx_path)

        for paragraph in doc.paragraphs:
            if "{{ imagem }}" in paragraph.text:
                # Remove apenas o texto do placeholder
                paragraph.clear()  # Limpa o conteúdo do parágrafo, mas mantém a posição

                # Adiciona a imagem dentro do mesmo parágrafo
                run = paragraph.add_run()
                run.add_picture(imagem_path, width=Inches(2))
                break


        output_path = os.path.join(os.path.dirname(docx_path), "saida.docx")
        doc.save(output_path)
        messagebox.showinfo("Sucesso", f"Arquivo salvo como:\n{output_path}")

    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro:\n{str(e)}")

# Interface gráfica com tkinter
class App:
    def __init__(self, master):
        self.master = master
        master.title("Inserir Imagem no Word")

        self.docx_path = None
        self.imagem_path = None

        self.label = tk.Label(master, text="Substituir {{ imagem }} por imagem em um .docx")
        self.label.pack(pady=10)

        self.btn_template = tk.Button(master, text="Selecionar Template (.docx)", command=self.selecionar_template)
        self.btn_template.pack(pady=5)

        self.btn_imagem = tk.Button(master, text="Selecionar Imagem", command=self.selecionar_imagem)
        self.btn_imagem.pack(pady=5)

        self.btn_processar = tk.Button(master, text="Processar e Salvar", command=self.processar)
        self.btn_processar.pack(pady=20)

    def selecionar_template(self):
        self.docx_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.docx_path:
            messagebox.showinfo("Arquivo selecionado", f"Template: {self.docx_path}")

    def selecionar_imagem(self):
        self.imagem_path = filedialog.askopenfilename(filetypes=[("Imagens", "*.png *.jpg *.jpeg")])
        if self.imagem_path:
            messagebox.showinfo("Imagem selecionada", f"Imagem: {self.imagem_path}")

    def processar(self):
        if not self.docx_path or not self.imagem_path:
            messagebox.showwarning("Atenção", "Selecione um arquivo .docx e uma imagem primeiro.")
        else:
            substituir_placeholder_por_imagem(self.docx_path, self.imagem_path)

# Executa a aplicação
if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.geometry("400x250")
    root.mainloop()
