from openpyxl import load_workbook
from docx import Document
import tkinter as tk
from datetime import datetime


doc = Document()


lista_botoes = []
# Carregar o arquivo excel
arquivo = load_workbook('OpenPy.xlsx')
planilha = arquivo.active

# Pegando os valores da coluna A para a lista, excluindo linhas que estão vazias
lista_botoes = [cell.value for cell in planilha['A'] if cell.value is not None]
print('Array:', lista_botoes)

# Function to run when a button is clicked
def on_click(botao_texto):
    print(f"Botão clicado: {botao_texto}")
    # Here you could add logic to write to Word, etc.


data = planilha['B2'].value
data_str = data.strftime('%Y-%m-%d')
random_data = planilha['B4'].value


def gerar_relatorio():
    #doc = Document()
    doc.add_heading(planilha['B1'].value, level= 1)
    doc.add_paragraph(data_str)
    doc.add_paragraph(random_data)
    doc.save("relatório.docx")
    print("Relatório salvo com sucesso!")
    print(type(planilha['B1'].value))


def adicionar_cabecalho(variavel):
    doc.add_heading(variavel)
    


# Setup the GUI window
root = tk.Tk()
root.geometry("500x300")
root.title('Organizador de Relatório')

# Adiciona botões em um layout de grade
cols = 3  # Número de botões por linha
for i, texto in enumerate(lista_botoes):
    row = i // cols
    col = i % cols
    botao = tk.Button(root, text=texto, command=lambda t=texto: on_click(t))
    botao.grid(row=row, column=col, padx=10, pady=10)

total_rows = (len(lista_botoes) + cols -1)
botao_gerar = tk.Button(root, text="Gerar relatório", command=gerar_relatorio)
botao_gerar.grid(row=total_rows, column=0, columnspan=cols, pady=20)
botao_gerar.pack

root.mainloop()

# Save the Excel file
arquivo.save("OpenPy")





# 1 - Ler coluna A1 ☑
# 2 - Separar cada item da coluna em um array ☑
# 3 - Cada item do array deve ser um botão clicável ☑
# 4 - Cada botão deve ter a função de ser adicionado a uma página word para impressão
# 5 - Converter datetime para data abreviada